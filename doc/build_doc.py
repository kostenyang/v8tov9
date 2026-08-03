# -*- coding: utf-8 -*-
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

SHOTS = r"C:\Users\mjalan\Documents\vspher8to9\shots"
OUT   = r"C:\Users\mjalan\Documents\vspher8to9\doc\VCF-5.2.1-to-9.1-Manual-Upgrade.docx"

d = Document()
st = d.styles['Normal']; st.font.name = 'Microsoft JhengHei'; st.font.size = Pt(10.5)

def H(t, lvl=1): return d.add_heading(t, level=lvl)
def P(t, bold=False, color=None, size=None, italic=False):
    p = d.add_paragraph(); r = p.add_run(t); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = RGBColor(*color)
    if size: r.font.size = Pt(size)
    return p
def bullet(t, bold=False):
    p = d.add_paragraph(style='List Bullet'); r = p.add_run(t); r.bold = bold; return p
def code(t):
    p = d.add_paragraph(); r = p.add_run(t); r.font.name='Consolas'; r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
    p.paragraph_format.left_indent=Inches(0.3); return p
def table(headers, rows, widths=None):
    t = d.add_table(rows=1, cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; run=c.paragraphs[0].add_run(h); run.bold=True; run.font.size=Pt(9.5)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; run=cells[i].paragraphs[0].add_run(str(v)); run.font.size=Pt(9.5)
    if widths:
        for i,w in enumerate(widths):
            for r in t.rows: r.cells[i].width=Inches(w)
    return t
def figure(fn, cap, w=6.0):
    img=os.path.join(SHOTS, fn)
    if os.path.exists(img):
        d.add_paragraph()
        pic=d.add_paragraph(); pic.alignment=WD_ALIGN_PARAGRAPH.CENTER
        pic.add_run().add_picture(img, width=Inches(w))
        c=d.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=c.add_run(cap); r.font.size=Pt(8.5); r.italic=True; r.font.color.rgb=RGBColor(0x66,0x66,0x66)

# ---------- Title ----------
tp=d.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=tp.add_run("VMware Cloud Foundation 5.2.1 → 9.1\n手動逐元件升級 交付文件"); r.bold=True; r.font.size=Pt(20)
sp=d.add_paragraph(); sp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=sp.add_run("Nested Lab (kosten.lab) · 非支援 Broadwell CPU · 全程實測記錄"); r.font.size=Pt(11); r.font.color.rgb=RGBColor(0x66,0x66,0x66)
mp=d.add_paragraph(); mp.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=mp.add_run("更新日期 2026-07-13 · Lab / 研究用途 · Broadcom 不支援此組態"); r.font.size=Pt(9); r.font.color.rgb=RGBColor(0x99,0x99,0x99)
d.add_paragraph()

# ---------- 1 摘要 ----------
H("1. 執行摘要與最終結果", 1)
P("在既有 nested VCF 5.2.1 環境上，脫離 SDDC Manager LCM、以手動 + CLI/API/UI 方式逐元件升級到 9.1，"
  "並在 Broadcom 已列為 discontinued 的 Broadwell CPU 上完成。全部元件已升級完成、環境健康。"
  "過程中最關鍵的技術突破，是解決了『nested 環境 vCLS 叢集服務 VM 無法開機』導致 DRS 失效、"
  "NSX Host 升級卡關的根因 (詳見第 6 節)；另一項重要發現是 profile-update 路徑升級 ESXi 9.1 "
  "在 Broadwell 上「不需要」allowLegacyCPU (第 9 節)。", bold=True)
table(
    ["元件", "版本 (前 → 後)", "最終狀態"],
    [
        ["NSX (整體)", "4.2.1 → 9.1.0.0100.25470810", "✅ 完成"],
        ["├ NSX Manager (MP)", "4.2.1 → 9.1", "✅ SUCCESS"],
        ["├ NSX Host VIB × 4", "4.2.1 → 9.1", "✅ SUCCESS (4/4)"],
        ["├ NSX Edge", "—", "✅ SUCCESS (無 Edge)"],
        ["└ NSX Finalize (Part 2)", "—", "✅ SUCCESS (vCenter 9.1 後解除 gate)"],
        ["vCenter", "8.0.3.00300 → 9.1.0.0100 (b25417926)", "✅ 完成 (migration，切回 .142)"],
        ["ESXi × 4", "8.0.3 → 9.1.0 (build-25433460)", "✅ 完成 (全 CONNECTED)"],
        ["vSAN", "—", "✅ OVERALL green"],
        ["SDDC Manager", "5.2.1", "⏭ 手動測試略過 (VCF LCM 專用)"],
    ],
    widths=[2.3, 2.6, 2.8]
)
P("殘留非阻擋警告：vSAN『Disk format version』(on-disk 格式可選升級到 9.1 格式)、『Performance service』(未啟用)。", size=9)

# ---------- 2 環境 ----------
H("2. 環境", 1)
table(
    ["角色", "版本 (升級前)", "IP"],
    [
        ["實體 ESXi (外層)", "8.0.3 build-24022510 · Xeon E5-2682 v4 (Broadwell-EP)", "192.168.110.16"],
        ["外層 vCenter (Xindian)", "8.0.3", "192.168.110.32"],
        ["Nested vCenter (VCF 管理域)", "8.0.3.00300", "192.168.110.142"],
        ["Nested NSX Manager", "4.2.1.0.0", "192.168.110.143"],
        ["Nested ESXi × 4", "8.0.3", "192.168.110.145–148"],
        ["SDDC Manager", "5.2.1", "192.168.110.141"],
        ["Jump host (repo/installer)", "Windows + PowerCLI", "192.168.110.200"],
    ],
    widths=[2.6, 3.1, 2.0]
)
P("升級媒體 (E:\\)：VMware-VCSA-all-9.1.0.0100.25417926.iso、VMware-ESXi-9.1.0.0100.25433460-depot.zip、"
  "VMware-NSX-upgrade-bundle-9.1.0.0100.0.25470810.mub", size=9)

# ---------- 3 官方順序 ----------
H("3. 官方升級順序 (VCF 9.1)", 1)
P("依 Broadcom KB 440630 (VCF 9.1) / KB 390634 (9.0) 與官方 techdocs：", bold=True)
P("SDDC Manager → NSX Manager (Part 1) → vCenter → ESXi → NSX Edge + Finalize (Part 2)", bold=True, size=12)
bullet("官方明確點出：NSX 9.1 升級『拆成兩部分』——Part 1 為 NSX Manager (vCenter 之前)；Part 2 為 NSX Edge + Finalize，且在 vCenter、ESX 之後才做。")
bullet("NSX Finalize 硬性要求 vCenter 已是 9.1 (本案實測 error 36013，見第 7 節)。")
bullet("本案對照：走傳統 NSX Upgrade Coordinator (脫離 LCM)，其精靈為 Manager→Host→Edge→Finalize 單一流程；NSX 4.2.1→9.1 的 host VIB 允許跑在 vCenter 8.0.3 上 (過渡)，故本案先完成 NSX Manager+Host，再升 vCenter，最後 ESXi 與 NSX Finalize。")

# ---------- 4 NSX Manager ----------
H("4. NSX 4.2.1 → 9.1：Manager (MP) 升級", 1)
P("以 URL 匯入 .mub (5.57GB) → 接受 EULA → 升級 Upgrade Coordinator → precheck → 升 NSX Manager (MP)。")
bullet("UI 卡點：NSX 9.1 Upgrade Coordinator UI (Angular19 iframe) 在 Chrome 語系為 zh-TW 時整片空白 (assets/i18n/zh-TW.json 回 HTTP 500)。解法：Chrome 語言改英文即可正常渲染。")
bullet("MP 升級一度在 98% 的 mp_notify_completion 失敗 (error 26：compute managers unhealthy)。根因：compute manager (vCenter) connection=DOWN。解法：PUT /api/v1/fabric/compute-managers/{id} 重新註冊 (帶帳密+thumbprint) → connection UP → retry → MP=SUCCESS。")
figure("10-nsx-mgr-success.png", "圖 4-1　NSX Manager (MP) 升級成功，5 個步驟全綠")

# ---------- 5 NSX Host ----------
H("5. NSX 4.2.1 → 9.1：Host VIB 升級 (VLCM 叢集)", 1)
P("本叢集為 vSphere Lifecycle Manager (vLCM) 影像管理，NSX Host 升級 = 交由 vLCM 對叢集影像加入 NSX 元件後逐台 remediate，"
  "每台 host 需進入 maintenance mode。第一次啟動即失敗：")
P("[vcf-m01-cl01] Host 'vcf-m01-esxNN' was not processed, reason: 'Health Check for cluster failed'", italic=True, color=(0xB0,0,0), size=9.5)
figure("13-nsx-hosts-failed.png", "圖 5-1　NSX Host 升級第一次失敗：叢集 Health Check 未通過，4 台皆未處理")
P("以 vLCM remediation precheck API 追出兩個 nested-lab 造成的阻擋 (POST /api/esx/settings/clusters/{id}/software?action=check)：", bold=True)
table(
    ["阻擋檢查", "根因", "解法"],
    [
        ["Host in vSAN maintenance mode possible (ERROR，全 4 台)",
         "vSAN 健康測試 controllerdiskmode『控制器磁碟群組模式為 VMware 認證』呈黃色 (nested = 未認證虛擬控制器)；vLCM 以 overall_health_not_green 擋 MM",
         "以 VsanHealthSetVsanClusterSilentChecks 靜音 HCL/控制器類健康測試 (短 id：controllerdiskmode / controlleronhcl / hcldbuptodate…)"],
        ["Is host ready to enter maintenance (ERROR，esx01/esx02)",
         "DRS is non-functional due to absence of vCLS quorum (Active vCLS VMs:0)。管理 VM 所在 host 無法自動 evacuate",
         "修復 vCLS (第 6 節) → DRS 恢復 → 通過"],
    ],
    widths=[2.4, 3.0, 2.0]
)

# ---------- 6 vCLS 突破 ----------
d.add_page_break()
H("6. 關鍵突破：nested vCLS 無法開機 → DRS 失效的根因與解法", 1)
P("這是本次升級最耗時、也最有價值的發現。NSX Host 升級需要 host 進 maintenance mode → 需要 DRS 自動搬移管理 VM → "
  "需要 vSphere Cluster Services (vCLS) 有 quorum。但本環境 3 個 vCLS VM 全部無法開機、反覆重建 (churn)，DRS 因此失效。", bold=True)
P("逐層追查：", bold=True)
bullet("vCLS VM 全數 POWERED_OFF，手動開機被拒 (403，service-managed)；retreat mode 切換、hostd 重啟、/etc/vmware/config 加 vmx.autoAnswer 皆無效。")
bullet("vCLS 為 CRX 微型 VM (路徑 /var/run/crx/infra/…，datastore 為空)。vmkernel.log 顯示其 vmm 世界每 ~0.3 秒重啟一次 = guest 反覆 reset。")
bullet("關鍵證據在 CRX 的 vmware.log：")
code("Power on failure messages: Feature 'cpuid.mwait' was 0, but must be 0x1.\nModule 'FeatureCompatLate' power on failed.")
P("根因 (對應 William Lam 文章)：vSphere 8 的 embedded vCLS (CRX) 需要 MONITOR/MWAIT 指令，"
  "而該指令『預設只對兩種 guest OS type 開啟：VMware ESXi 與 macOS』。本環境的 nested ESXi VM 之 guest OS type 被建為 "
  "『Other (64-bit) / OTHER_64』→ 不開放 MWAIT → CRX vCLS 無法開機。", bold=True, color=(0xB0,0,0))
P("解法 (在實體/外層 Xindian vCenter 192.168.110.32 上，對 4 台 nested ESXi VM 逐台)：", bold=True)
bullet("① 升級虛擬硬體版本 vmx-14 → vmx-20 (原 HW14=ESXi6.7 等級，導致 vmkernel8Guest 被判 UnsupportedGuest，即使實體 host 為 8.0.3)。")
bullet("② 將 guest OS type 改為 VMware ESXi 8.x (guestId=vmkernel8Guest)。")
bullet("③ 完整關機再開機 (滾動逐台，vSAN FTT=1 可容忍單台離線；管理 VM 先手動/DRS vMotion 移走)。")
P("驗證：完成 esx03、esx04 兩台後，vCLS VM 即成功開機、DRS 恢復 functional、vLCM precheck 全數通過，NSX Host 升級隨即得以進行。"
  "(僅需 1 個 vCLS quorum；此修復存於 VM vmx 為持久設定 — 事後一次實體停電，環境恢復後 vCLS 自動開機、vSAN 資料完整。)", bold=True)

# ---------- 7 NSX host success + finalize ----------
H("7. NSX Host 升級完成 與 Finalize gate", 1)
P("以 session+XSRF API 續跑 (UI 之 UC iframe 常凍結)：POST /api/v1/upgrade/plan?action=continue&component_type=HOST，"
  "隨即 IN_PROGRESS 並越過先前失敗的 health check。4 台以 serial 逐台完成 (esx04→esx02→esx01→esx03)，"
  "DRS 自動搬移含 vCenter/NSX/SDDC 在內的管理 VM，全程約 30 分鐘。")
figure("16-nsx-hosts-success.png", "圖 7-1　NSX Host 升級 Upgrade Status = Success (4/4，No HOST(s) with issues)")
figure("17-nsx-finalize.png", "圖 7-2　NSX Finalize Upgrade 5 步驟 (Creating local backup / Notify Completion / Product Version Update / Unpin API …)")
P("NSX Finalize 被 gate 擋 (符合官方順序 — Finalize 屬 Part 2，需 vCenter 9.1)：", bold=True)
P("error 36013：Found compute manager at version 8.0.3. Please upgrade the compute manager to version 9.1 to complete the NSX upgrade.",
  italic=True, color=(0xB0,0,0), size=9.5)
P("→ 必須先把 vCenter 升到 9.1，才能執行 NSX Finalize。(EDGE 無節點，以 API action=upgrade&component_type=EDGE 標記為 SUCCESS。)")
P("★ 最終結果：vCenter + ESXi 升級完成後，以 API 觸發 FINALIZE_UPGRADE → 約 3 分鐘後 status=SUCCESS，"
  "NSX 版本標記為 9.1.0.0100.25470810，MP / HOST / EDGE / FINALIZE 全部 SUCCESS。", bold=True, color=(0,0x70,0))

# ---------- 8 vCenter ----------
H("8. vCenter 8.0.3 → 9.1 (CLI migration)", 1)
P("VAMI『更新』分頁不適用 (8→9 為大版本升級、非 patch；且 VAMI 自訂 repo 只吃 HTTPS)。正解為 VCSA 9.1 ISO 的 CLI migration installer。")
code('F:\\vcsa-cli-installer\\win32\\vcsa-deploy.exe upgrade --accept-eula --acknowledge-ceip\n'
     '  --no-ssl-certificate-verification --upgrade-framework legacy\n'
     '  --skip-product-interop-check --verbose  vcsa91-upgrade.json')
P("依序踩到的坑與解法：", bold=True)
bullet("RDU 卡死：預設 reduced_downtime_upgrade 在 nested DRS 叢集資源驗證卡住 → 改 --upgrade-framework legacy。")
bullet("NSX 相容性硬擋：來源 vCenter 的 upgrade checker 曾回『NSX 4.2.1 not supported』→ 先升 NSX (本案已 9.1) 後此擋自動解除。")
bullet("新 9.1 appliance 為 Linux(Photon) guest VM，不受 Broadwell CPU 封鎖 (CPU 封鎖只在 ESXi hypervisor 層)。")
bullet("migration 需暫時 IP (.151，部署新機用，遷移後切回原 IP .142)；vcdb.migrateSet=core。")
P("✅ precheck 全數通過 (Host requirements are met with DRS enabled — 即 vCLS 修復生效)；僅 2 個非阻擋警告 "
  "(LCM 不相容檔案不複製、外部 extension 需升級後重新註冊)。", bold=True)
P("★ 升級完成 (約 52 分鐘)：新 appliance 部署到 temp IP .151 → 子任務 rpminstall / precheck / export / firstboot / import 全數 SUCCEEDED "
  "→ 自動切回原 IP .142、舊機關機。驗證：vCenter 9.1.0.0100 build 25417926、SSO 正常、4 host CONNECTED、inventory 完整。",
  bold=True, color=(0,0x70,0))
bullet("坑：升級中 installer 曾拋一次 SecureConnectorException (取新機 5480 憑證回 None，訊息暗示 clock skew 但時鐘正常 ±7 秒) — "
       "屬一次性 glitch，installer 自行重試恢復，非致命。權威進度來源是新機的 GET https://<新IP>:5480/rest/vcenter/deployment (root basic auth)。")

# ---------- 9 ESXi ----------
H("9. ESXi 8.0.3 → 9.1 在 Broadwell CPU 上 (4 台全部完成)", 1)
P("Xeon E5-2682 v4 = Broadwell-EP，屬 Broadcom KB 318697 discontinued (ISO 安裝程式封鎖)。實測流程 (4 台皆成功，build 25433460)：")
bullet("① depot.zip 解壓為 online depot (含 index.xml)，以 HTTP 掛出 (esxcli offline -d 只吃本機路徑，不吃 URL)。")
bullet("② host 進 maintenance mode (vSAN ensureObjectAccessibility，DRS 自動搬走管理 VM)。")
bullet("③ esxcli software profile update -d <depot>/index.xml -p ESXi-9.1.0.0100-25433460-standard --no-hardware-warning")
bullet("④ reboot → 退出 maintenance mode → 等 vSAN resync 歸零，再升下一台。")
P("★ 重大發現：透過 profile-update 路徑升級，Broadwell 上「不需要」allowLegacyCPU 即可開機 9.1。", bold=True, color=(0,0x70,0))
P("原本依既有資料預期需在 boot.cfg 的 kernelopt 加 allowLegacyCPU=TRUE。實測第一台 (esx03) 因 sed 跳脫問題該參數其實「沒有」寫入 "
  "(boot.cfg 仍為 kernelopt=autoPartition=FALSE)，但 host 仍成功開機為 ESXi 9.1.0 build-25433460。後續 3 台遂直接略過該步驟，"
  "4 台全部順利升級。結論：KB 318697 的 CPU 封鎖是「ISO 安裝程式」層的檢查；以 esxcli profile update + --no-hardware-warning "
  "的升級路徑不會在開機時硬擋 Broadwell。(allowLegacyCPU 仍可作為保險，但本案實證非必要。)", bold=True)
P("升級順序：先升沒有 running 管理 VM 的 host；跑 vCenter 的 host 最後升 (MM 時 DRS 會 vMotion vCenter 自己，API 短暫抖動)。", size=9)

# ---------- 10 卡點速查 ----------
d.add_page_break()
H("10. 卡點速查表", 1)
table(
    ["症狀", "原因 / 解法"],
    [
        ["nested vCLS 無法開機 / DRS 失效 / cpuid.mwait was 0", "nested ESXi VM guest type=Other(64bit) 不開 MWAIT → 改 guestId=vmkernel8Guest + 升 vHW-20 (外層 vCenter，需關機)"],
        ["NSX Host 升級 Health Check for cluster failed", "vLCM MM precheck 擋：① vSAN controllerdiskmode 黃 → 靜音 HCL 健康測試；② vCLS/DRS quorum → 修 vCLS"],
        ["NSX Finalize error 36013 compute manager 8.0.3", "先升 vCenter 到 9.1 才能 finalize"],
        ["NSX UC UI 整片空白", "Chrome 語系 zh-TW 的 i18n JSON 回 500 → 改英文"],
        ["NSX UC UI 凍結 (CDP screenshot timeout)", "改用 session+XSRF API 驅動 (action=continue/upgrade&component_type=…)"],
        ["ESXi 9 裝不上 (ISO 安裝程式) / Unsupported CPU", "Broadwell discontinued。改走 esxcli profile update + --no-hardware-warning；★實測此路徑「不需」allowLegacyCPU"],
        ["vcsa-deploy RDU DRS resource validation 卡死", "--upgrade-framework legacy"],
        ["vcsa-deploy NSX 4.2.1 not supported", "先升 NSX (官方順序)"],
        ["vcsa-deploy SecureConnectorException (NoneType / 疑 clock skew)", "一次性 glitch，會自行重試恢復；用新機 :5480/rest/vcenter/deployment 看權威進度"],
        ["VAMI 更新頁看不到 9.1 / repo URL 紅字", "8→9 非 patch，改 ISO CLI migration；VAMI 自訂 repo 只吃 HTTPS"],
        ["plink 送 sed 到 ESXi 跳脫字元壞掉", "避免 $ / 引號巢狀；或改用不含特殊字元的命令 (本案最後證實該步驟非必要)"],
    ],
    widths=[3.1, 4.2]
)

# ---------- 11 最終結果 ----------
H("11. 最終結果與收尾", 1)
P("★ 升級全數完成 (2026-07-13)，環境健康：", bold=True, color=(0,0x70,0))
table(
    ["元件", "最終版本 / 狀態"],
    [
        ["NSX", "9.1.0.0100.25470810 — MP / HOST / EDGE / FINALIZE 全 SUCCESS"],
        ["vCenter", "9.1.0.0100 build 25417926 (原 IP .142，SSO 正常，inventory 完整)"],
        ["ESXi × 4", "9.1.0 build-25433460 — 全 CONNECTED、maintenance mode Disabled"],
        ["vSAN", "OVERALL green — 物件全 healthy、resync 0"],
    ],
    widths=[2.0, 5.3]
)
P("殘留 (非阻擋，可選處理)：", bold=True)
bullet("vSAN『Disk format version』黃燈：on-disk 格式可選升級到 vSAN 9.1 格式 (需逐台 disk group 升級)。")
bullet("vSAN『Performance service』黃燈：效能服務未啟用 (本 lab 一直如此)。")
bullet("外部 extension (vSphere LCM Client / SDDC Manager Plugin / NSX Manager) 建議於 vCenter 升級後確認/重新註冊。")
bullet("SDDC Manager 仍為 5.2.1 (本案刻意脫離 VCF LCM；如需完整 VCF 9.1 fleet 需另走 LCM/Installer 路線)。")

P("")
P("※ 本文件為 Lab / 研究用途實測記錄。文中使用的旁路 (allowLegacyCPU、guestId 改 ESXi type、靜音 vSAN HCL 健康測試、"
  "--skip-product-interop-check 等) 皆會使環境脫離 Broadcom 受支援狀態；正式環境請走 SDDC Manager LCM 並使用受支援硬體。",
  size=9, color=(0x99,0x99,0x99), italic=True)

d.save(OUT)
print("SAVED:", OUT)
