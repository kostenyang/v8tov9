# -*- coding: utf-8 -*-
# build_doc_rebuild.py — VCF 5.2.1 → 9.1 完整重建 + Ops/License 部署 交付文件
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = r"C:\Users\mjalan\Documents\vspher8to9"
SHOTS = os.path.join(ROOT, "shots")
OUT = os.environ.get("DOCX_OUT", os.path.join(ROOT, "doc", "VCF-5.2.1-to-9.1-Rebuild-Report.docx"))

d = Document()
# 基本樣式
st = d.styles['Normal']; st.font.name='Calibri'; st.font.size=Pt(10.5)
st.element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')

def H(t,lvl=1):
    h=d.add_heading(t,level=lvl)
    for r in h.runs:
        r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    return h
def P(t,bold=False,italic=False,size=10.5,color=None):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=italic; r.font.size=Pt(size)
    r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    if color: r.font.color.rgb=color
    return p
def BULLET(t):
    p=d.add_paragraph(style='List Bullet'); r=p.add_run(t)
    r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei'); r.font.size=Pt(10.5)
    return p
def FIG(fname,caption,width=6.3):
    fp=os.path.join(SHOTS,fname)
    if os.path.exists(fp):
        d.add_picture(fp,width=Inches(width))
        d.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
        c=d.add_paragraph(); r=c.add_run("圖："+caption); r.italic=True; r.font.size=Pt(9)
        r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
        c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    else:
        P("[缺圖 "+fname+"]",italic=True,color=RGBColor(0xC0,0,0))
def TABLE(headers,rows):
    t=d.add_table(rows=1,cols=len(headers)); t.style='Light Grid Accent 1'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=''; r=c.paragraphs[0].add_run(h); r.bold=True; r.font.size=Pt(9.5)
        r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row):
            cells[i].text=''; r=cells[i].paragraphs[0].add_run(str(v)); r.font.size=Pt(9.5)
            r.font.name='Calibri'; r._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
    return t

# ===== 封面 =====
title=d.add_heading('',level=0)
rt=title.add_run('VMware Cloud Foundation 5.2.1 → 9.1\n完整重建與升級交付報告')
rt.font.name='Calibri'; rt._element.rPr.rFonts.set(qn('w:eastAsia'),'Microsoft JhengHei')
title.alignment=WD_ALIGN_PARAGRAPH.CENTER
P('Nested Lab（Broadwell CPU）· 手動升級路徑（脫離 SDDC Manager LCM）· 含 VCF Operations + License Server 部署',italic=True,size=11).alignment=WD_ALIGN_PARAGRAPH.CENTER
P('報告日期：2026-08-04',size=10).alignment=WD_ALIGN_PARAGRAPH.CENTER
d.add_paragraph()

# ===== 1 摘要 =====
H('1. 執行摘要',1)
P('本報告記錄在巢狀實驗環境（4 台 nested ESXi + vSAN，實體 CPU 為 Intel Xeon E5-2682 v4 Broadwell，於 VCF 9.x 已列 discontinued）上，從零重建 VCF 5.2.1 管理域，再以手動方式（不透過 SDDC Manager LCM）逐一升級至 9.1，最後部署 VCF Operations 與 License Server 的完整過程、關鍵技術障礙與解法、以及各階段實機截圖。')
P('最終成果：',bold=True)
TABLE(['元件','版本（升級後）'],[
 ['NSX','9.1.0.0200.25524170（MP / Host×4 / Finalize 全 SUCCESS）'],
 ['vCenter Server','9.1.0 build 25629530（migration-based）'],
 ['ESXi ×4','9.1.0 build 25433460（Broadwell，profile-install 路徑）'],
 ['vSAN','健康，resync=0'],
 ['VCF Operations','9.1，單節點叢集 INITIALIZED'],
 ['VCF License Server','已部署開機（.151）；正式授權註冊需 Broadcom BSC（見 §8）'],
])

# ===== 2 環境 =====
H('2. 環境與位址',1)
TABLE(['角色','FQDN / 名稱','IP','備註'],[
 ['外層 Xindian vCenter','—','192.168.110.32','部署 nested VM 的實體管理層'],
 ['Nested vCenter','vcf-m01-vc01','192.168.110.142','8.0.3 → 9.1'],
 ['SDDC Manager','vcf-m01-sddcm01','192.168.110.141','5.2.1（刻意脫離 LCM）'],
 ['NSX Manager','vcf-m01-nsx01','192.168.110.143','4.2.1 → 9.1.0.0200'],
 ['Nested ESXi ×4','vcf-m01-esx01~04','192.168.110.145–148','8.0.3 → 9.1'],
 ['VCF Operations','vcf-ops01','192.168.110.150','9.1（新部署）'],
 ['VCF License Server','vcf-lic01','192.168.110.151','9.1（新部署）'],
 ['Cloud Builder','vcf-m01-cb01','192.168.110.140','5.2.1 bringup 控制器'],
 ['AD / DNS','—','192.168.110.1','kosten.lab；GW .254'],
])

# ===== 3 5.2.1 重建 =====
H('3. Phase 1–4：重建 VCF 5.2.1 管理域',1)
P('毀掉既有環境後，以 PowerCLI 於 Xindian vCenter 部署 4 台 nested ESXi 8.0u3（各 8 vCPU / 64GB / vSAN cache+capacity）+ Cloud Builder 5.2.1，產生 bringup 設定檔（License-Later 評估模式），再由 Cloud Builder 執行 142 個子任務完成管理域 bringup。')
BULLET('部署腳本：deploy-vcf521-m01.ps1（nested ESXi + Cloud Builder）、run-bringup.ps1（validate → deploy → monitor）')
BULLET('前置：ADSrv 上為所有 FQDN 建立正/反解 DNS（prep-dns.ps1）')
BULLET('bringup 結果：COMPLETED_WITH_SUCCESS（約 2 小時 40 分）')
FIG('p4-cloudbuilder-login.png','VMware Cloud Builder（5.2.1 bringup 控制器）')
FIG('p4-sddcm-521.png','bringup 完成 — SDDC Manager / vSphere SSO 登入（Managed by SDDC Manager）')
FIG('p4-nsx-521.png','bringup 完成 — NSX Manager 4.2.1 登入頁（升級前基準）')

# ===== 4 NSX 升級 =====
H('4. Phase 5a：NSX 4.2.1 → 9.1 升級',1)
P('依官方順序（KB 440630）NSX 先升。流程：接受 EULA → 匯入 9.1 升級 bundle（.mub，經本機 HTTP 由 URL 匯入）→ 升級 Upgrade Coordinator → 啟動計畫（pre-check + MP）→ Host（vLCM VIB 滾動）→ Edge（無節點 skip）→ Finalize（需 vCenter 9.1）。全程以 session + XSRF token 的 REST API 驅動。')
P('關鍵技術障礙與解法：',bold=True)
BULLET('EULA 需「重接受」才生效（首次 acceptance=false）；bundle 上傳狀態 endpoint 為 /upload-status；啟動計畫用 POST /api/v1/upgrade/plan?action=start。')
BULLET('Host 升級前必備修復：nested ESXi VM 原為 otherGuest64 / vmx-14 → MWAIT 關閉 → vCLS（CRX 微 VM）起不來 → DRS 失效 → vLCM 擋 host 進 MM。修法：於 Xindian 將 esx03/esx04 兩台改 guestId=vmkernel8Guest + vHW-20（upgrade-hw-guest.ps1），vCLS 即在其上開機取得 quorum、DRS 恢復。')
BULLET('靜音 vSAN HCL 健康測試（vsan-silence.ps1），否則 nested 虛擬控制器未認證使 vLCM 以 overall_health_not_green 擋 host MM。')
BULLET('Finalize 不會被 continue 推進，需 POST /api/v1/upgrade/plan?action=upgrade&component_type=FINALIZE_UPGRADE。')
P('以下為 NSX Upgrade Coordinator（System → Upgrade）的逐步實機畫面：',bold=True)
FIG('up01-nsx-before-421.png','① 升級前：NSX 4.2.1，Upgrade Coordinator 起點（Prepare for Upgrade）',5.5)
FIG('up05-mp-done.png','② UPGRADE SUMMARY：Prepare ✅、Upgrade NSX Manager ✅（MP 完成，Target 9.1.0.0200.25524170）',5.5)
FIG('up06-host-progress.png','③ Upgrade Hosts 進行中（右上 “In Progress” 徽章）',5.5)
FIG('up07-host-done.png','④ Upgrade Hosts ✅（MP + Hosts 完成，Edges/Finalize 待續）',5.0)
FIG('up11-all-done.png','⑤ 升級完成：“Complete” 徽章，Check Readiness / Prepare / NSX Manager / Hosts / Edges & VNAs / Finalize 全 ✅，版本 9.1.0.0200.25524170',5.0)

# ===== 5 vCenter 升級 =====
H('5. Phase 5b：vCenter 8.0.3 → 9.1 Migration',1)
P('以 vcsa-deploy CLI（VCSA 9.1 ISO）執行 migration-based 升級：部署新 9.1 appliance 至暫時 IP .151，遷移組態後自動切換回 .142、關閉舊機。')
BULLET('旗標：--upgrade-framework legacy（預設 RDU 在 nested DRS 資源驗證卡死）、--skip-product-interop-check；範本 vcsa91-upgrade.json（vcdb_migrateSet=core）。')
BULLET('precheck exit 0 → 正式 migration 約 50 分完成，所有子任務 SUCCEEDED。')
P('vCenter migration 逐步（vcsa-deploy 各階段實際執行結果）：',bold=True)
TABLE(['#','階段 / Task','說明','結果'],[
 ['1','Stage 1：部署新 9.1 appliance','OVF Transfer → Power On，暫時 IP 192.168.110.151','完成'],
 ['2','Install required RPMs','在新 appliance 安裝所需 RPM','SUCCEEDED 100/100'],
 ['3','Run appliance precheck','升級前檢查','SUCCEEDED 100/100'],
 ['4','Export existing appliance data','從舊 8.0.3 匯出組態/資料','SUCCEEDED 100/100'],
 ['5','Run firstboot scripts','新 appliance firstboot 起各服務','SUCCEEDED 100/100'],
 ['6','Import data into appliance','把資料匯入新 9.1 appliance','SUCCEEDED 100/100'],
 ['7','IP 切換 + 收尾','.151→.142、舊機關機、評估模式','Successfully completed'],
])
FIG('up08-vcenter-vami-91.png','升級後（結果驗證）：vCenter VAMI Summary — Version 9.1.0.0300 / Build 25629530，Health 全 Good，SSO Running')

# ===== 6 ESXi 升級 =====
H('6. Phase 5c：ESXi ×4 → 9.1（Broadwell 繞過）',1)
P('順序：先升無執行管理 VM 的 host，跑 vCenter 的 host 最後。每台：進 MM（vSAN ensureObjectAccessibility，DRS 疏散）→ 由 vSAN datastore 上的本機 depot 執行升級 → reboot → 退 MM → 待 resync=0。')
P('關鍵技術障礙與解法：',bold=True)
BULLET('Broadwell CPU：profile-install/update 路徑開機 ESXi 9.1「不需」allowLegacyCPU（KB 318697 的 CPU 阻擋僅適用 ISO 安裝程式）。')
BULLET('depot 取得：nested host 連本機 HTTP:8080 被 Windows 防火牆擋（無 admin 權限加規則）→ 改將 692MB ESXi 9.1 depot 複製到 vSAN datastore，以本機路徑 -d /vmfs/volumes/.../esxi91-depot/index.xml 升級。')
BULLET('★ 0200 NSX bundle 特有衝突：host 上的 NSX VIB 為 9.1.0.0200-8.0 flavor（要求 esx-version << 8.1），使 profile update 到 ESXi 9.1 出現 DependencyError；live vib remove 又因 nsx-datapath 使用中失敗。解法：esxcli software profile install ... --ok-to-remove --no-live-install --maintenance-mode，一次移除 ESXi8.0 + NSX-8.0 VIB、裝上 ESXi 9.1 + NSX-9.1 VIB（depot 內含 nsx-*_9.1.0.0100-9.1）。vSAN disk 仍 In CMMDS，無需重加 mock VIB。')
BULLET('每台升級前以 vCenter 重新啟用 SSH（reboot 後預設關）。')
FIG('up09-esxi91-console.png','升級後：ESXi 主控台 DCUI — VMware ESXi 9.1.0.0100.25433460 on Intel Xeon E5-2682 v4（Broadwell，VCF9 列 discontinued，本 profile-install 法成功升級並開機）',5.5)
FIG('up10-esxi-hostclient-91.png','升級後：ESXi Host Client — Hypervisor VMware ESXi 9.1.0.0100.25433460，Processor Broadwell E5-2682 v4',5.5)
TABLE(['ESXi Host','升級後版本','vSAN'],[
 ['vcf-m01-esx01~04','9.1.0 build 25433460','healthy / resync=0'],
])

# ===== 7 Operations =====
H('7. Phase 6：VCF Operations 部署',1)
P('以 PowerCLI Import-VApp 部署 Operations Appliance（xsmall 2vCPU/8GB，thin）。')
P('關鍵技術點：',bold=True)
BULLET('ovftool 的 --prop 對本 OVA 的網路屬性（屬於 product instance VMware_Aria_Operations）無效 → 改用 PowerCLI Get-OvfConfiguration / Import-VApp 設定；部分字串屬性（root_password/domain）需於「關機狀態」以 ReconfigVM 補設。')
BULLET('firstboot 絕不可打斷 → 部署時不自動開機，關機設好 IP 後「單次開機」。')
BULLET('初始化：CaSA API（GET /casa/node/thumbprint → POST /casa/cluster init → 輪詢至 cluster_state=INITIALIZED）；appliance HTTPS 一律用 curl.exe（Windows PowerShell 5.1 的 .NET TLS 無法握手）。')
FIG('up12-ops-login.png','VCF Operations 登入頁（9.1，已上線）')
FIG('up12-ops-dashboard.png','登入後：VCF Operations 設定精靈「Congratulations on setting up VCF Operations」')

# ===== 8 License Server =====
H('8. Phase 6：VCF License Server 部署',1)
P('以 PowerCLI Import-VApp 部署 License Server Appliance（2vCPU/4GB，.151）。')
P('關鍵技術點與限制：',bold=True)
BULLET('OVA 簽章憑證 PowerCLI 不信任 → 解開 OVA、移除 .cert/.mf 成為 unsigned OVF 後再 Import-VApp。')
BULLET('網路屬性同屬 instance（VCF_License_Server_Appliance），以 Get-OvfConfiguration 設定；hostname/otk 為預設段屬性。')
BULLET('★ 授權限制（air-gapped）：VCF 9.1 的 License Server「唯一註冊金鑰（otk）」由 VCF Operations 產生，而其產生需先向 Broadcom Business Services Console（BSC，需網際網路）註冊；本無網路 lab 實測 /casa/license/registration-key 回 HTTP 500「Operation failed」。故本環境 License Server 完成「部署 + 開機」，實際授權註冊/分發需 BSC 連線（超出離線 lab 範圍）。')
FIG('up13-lic-console.png','VCF License Server 開機（.151）')

# ===== 9 附錄 =====
H('9. 附錄：關鍵腳本',1)
BULLET('deploy-vcf521-m01.ps1 / run-bringup.ps1 / prep-dns.ps1 — 5.2.1 部署與 bringup')
BULLET('upgrade-hw-guest.ps1 / vsan-silence.ps1 — Host 升級前 MWAIT/vCLS/vSAN 修復')
BULLET('nsx-precheck.ps1 + upgrade REST（session/XSRF）— NSX 升級驅動')
BULLET('vcsa91-upgrade.json — vCenter migration 範本')
BULLET('esxi-upgrade-host.ps1 — ESXi profile-install 升級（含 NSX VIB swap）')
BULLET('deploy-ops-pcli.ps1 / ops-setup2.ps1 — Operations 部署與 CaSA 初始化')
BULLET('deploy-lic-pcli.ps1 — License Server 部署')

d.save(OUT)
print("SAVED", OUT, os.path.getsize(OUT), "bytes")
